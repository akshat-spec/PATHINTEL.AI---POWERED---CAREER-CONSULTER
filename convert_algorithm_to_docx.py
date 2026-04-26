from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def create_algorithm_docx():
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # 1. Title
    title = doc.add_heading('Career Guidance and Job Fit Prediction Algorithm', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Comprehensive Technical Documentation')
    run.italic = True
    
    doc.add_paragraph()
    
    # 2. Executive Summary
    doc.add_heading('Executive Summary', level=1)
    p = doc.add_paragraph()
    p.add_run('This document provides a comprehensive explanation of the algorithms implemented in the Career Guidance System. The system uses multiple machine learning approaches to provide:')
    
    bullet1 = doc.add_paragraph()
    bullet1.add_run('Career Path Prediction - Based on resume analysis').bold = True
    
    bullet2 = doc.add_paragraph()
    bullet2.add_run('Job Fit Prediction - Match percentage between user profile and target job').bold = True
    
    bullet3 = doc.add_paragraph()
    bullet3.add_run('Skill Gap Analysis - Identification of missing skills and recommendations').bold = True
    
    doc.add_paragraph()
    
    # 3. Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc = [
        ('1. ', 'System Architecture'),
        ('2. ', 'Algorithm 1: Career Classification Model'),
        ('3. ', 'Algorithm 2: Job Fit Prediction Model'),
        ('4. ', 'Algorithm 3: Skill-Based Assessment Model'),
        ('5. ', 'Feature Engineering'),
        ('6. ', 'Data Processing Pipeline'),
        ('7. ', 'Model Training Process'),
        ('8. ', 'Prediction Workflow'),
        ('9. ', 'Ensemble Methods'),
        ('10. ', 'API Endpoints')
    ]
    
    for num, heading in toc:
        p = doc.add_paragraph(num + heading)
        p.paragraph_format.left_indent = Cm(1)
    
    doc.add_paragraph()
    
    # 4. System Architecture
    doc.add_heading('1. System Architecture', level=1)
    
    p = doc.add_paragraph()
    p.add_run('The Career Guidance System consists of three main components:')
    
    # Architecture components
    components = [
        ('CareerModel (career_model.py)', 'XGBoost classifier with Transformer embeddings for resume-based career prediction'),
        ('Skill Model (career_model.pkl)', 'Random Forest classifier for quiz-based skill assessment'),
        ('JobProbability Predictor (job_probability_model.py)', 'XGBoost regressor for job fit prediction')
    ]
    
    for i, (name, desc) in enumerate(components, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {name}: ').bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    
    # 5. Algorithm 1: Career Classification
    doc.add_heading('2. Algorithm 1: Career Classification Model', level=1)
    
    doc.add_heading('Purpose', level=2)
    p = doc.add_paragraph()
    p.add_run('Predict suitable career paths based on resume content analysis.')
    
    doc.add_heading('Processing Steps', level=2)
    
    # Step 1
    doc.add_heading('Step 1: Text Preprocessing', level=3)
    p = doc.add_paragraph()
    p.add_run('The resume text is cleaned by removing URLs, special characters (except for programming languages like C++, .NET), non-ASCII characters, and normalizing whitespace. The text is converted to lowercase for consistent processing.')
    
    # Step 2
    doc.add_heading('Step 2: Named Entity Recognition (NER)', level=3)
    p = doc.add_paragraph()
    p.add_run('Using a hybrid approach combining spaCy NER and keyword matching:')
    
    p = doc.add_paragraph('• Experience extraction using regex patterns')
    p = doc.add_paragraph('• Skills extraction from a 100+ skill database')
    p = doc.add_paragraph('• Education detection (B.Tech, M.Tech, PhD, etc.)')
    
    # Step 3
    doc.add_heading('Step 3: Feature Extraction', level=3)
    p = doc.add_paragraph()
    p.add_run('Using Sentence Transformers (all-MiniLM-L6-v2) to generate 384-dimensional embeddings of the cleaned resume text.')
    
    # Step 4
    doc.add_heading('Step 4: Classification', level=3)
    p = doc.add_paragraph()
    p.add_run('XGBoost Multi-class Classifier with the following parameters:')
    
    params = [
        'n_estimators=200',
        'learning_rate=0.05',
        'max_depth=6',
        'objective=\'multi:softprob\'',
        'random_state=42'
    ]
    
    for param in params:
        p = doc.add_paragraph(param)
        p.paragraph_format.left_indent = Cm(1)
    
    # Output
    doc.add_heading('Output', level=2)
    p = doc.add_paragraph()
    p.add_run('Returns top 5 career predictions with probability scores:')
    
    p = doc.add_paragraph()
    run = p.add_run('[{"role": "Data Scientist", "score": 85.5}, {"role": "Machine Learning Engineer", "score": 72.3}, ...]')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # 6. Algorithm 2: Job Fit Prediction
    doc.add_heading('3. Algorithm 2: Job Fit Prediction Model', level=1)
    
    doc.add_heading('Purpose', level=2)
    p = doc.add_paragraph()
    p.add_run('Calculate the match probability between a user\'s resume and a specific target job.')
    
    doc.add_heading('Feature Vector Composition', level=2)
    
    # Feature table
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    features = [
        ('Semantic Similarity', 'Float (0-1)', 'Cosine similarity between embeddings'),
        ('Skill Match Ratio', 'Float (0-1)', 'Overlap of skills / Total job skills'),
        ('Skill Overlap Count', 'Integer', 'Number of matching skills'),
        ('Experience Years', 'Integer', 'Years of experience extracted'),
        ('Keyword Density', 'Float', 'Skills density in resume'),
        ('Title Match Score', 'Float', 'Job title keyword overlap'),
        ('Bigram Overlap', 'Integer', 'Common bigrams between texts'),
        ('Resume Skills Count', 'Integer', 'Total skills found'),
        ('Job Skills Count', 'Integer', 'Total skills required')
    ]
    
    for i, (feature, ftype, desc) in enumerate(features):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = feature
        row_cells[1].text = ftype
        row_cells[2].text = desc
    
    doc.add_paragraph()
    
    # Confidence Classification
    doc.add_heading('Confidence Classification', level=2)
    
    conf_table = doc.add_table(rows=5, cols=3)
    conf_table.style = 'Table Grid'
    
    hdr = conf_table.rows[0].cells
    hdr[0].text = 'Probability Range'
    hdr[1].text = 'Confidence Level'
    hdr[2].text = 'Message'
    
    conf_data = [
        ('>= 80%', 'High', 'Excellent match!'),
        ('60-79%', 'Medium-High', 'Good match!'),
        ('40-59%', 'Medium', 'Decent match'),
        ('< 40%', 'Low', 'Significant gaps')
    ]
    
    for i, (prob, conf, msg) in enumerate(conf_data):
        row = conf_table.rows[i+1].cells
        row[0].text = prob
        row[1].text = conf
        row[2].text = msg
    
    doc.add_paragraph()
    
    # 7. Algorithm 3: Skill-Based Assessment
    doc.add_heading('4. Algorithm 3: Skill-Based Assessment Model', level=1)
    
    doc.add_heading('Purpose', level=2)
    p = doc.add_paragraph()
    p.add_run('Predict career based on user\'s self-assessed skill levels across 17 dimensions.')
    
    doc.add_heading('Input Encoding', level=2)
    
    enc_table = doc.add_table(rows=7, cols=2)
    enc_table.style = 'Table Grid'
    
    hdr = enc_table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Numeric Value'
    
    enc_data = [
        ('Not Interested', '0'),
        ('Poor', '1'),
        ('Beginner', '2'),
        ('Average', '3'),
        ('Intermediate', '4'),
        ('Excellent', '5')
    ]
    
    for i, (cat, val) in enumerate(enc_data):
        row = enc_table.rows[i+1].cells
        row[0].text = cat
        row[1].text = val
    
    doc.add_paragraph()
    
    # 8. Feature Engineering
    doc.add_heading('5. Feature Engineering', level=1)
    
    doc.add_heading('Skill Database Categories', level=2)
    
    categories = [
        ('Programming', 'python, java, javascript, c++, c#, ruby, php'),
        ('Web Development', 'html, css, react, angular, vue, node.js, django, flask'),
        ('Data Science', 'machine learning, deep learning, tensorflow, pytorch, pandas'),
        ('Databases', 'sql, mysql, postgresql, mongodb, redis'),
        ('Cloud & DevOps', 'aws, azure, gcp, docker, kubernetes'),
        ('Soft Skills', 'leadership, communication, problem solving, agile')
    ]
    
    for cat, skills in categories:
        p = doc.add_paragraph()
        p.add_run(f'{cat}: ').bold = True
        p.add_run(skills)
    
    doc.add_paragraph()
    
    # 9. Data Processing Pipeline
    doc.add_heading('6. Data Processing Pipeline', level=1)
    
    steps = [
        ('1. Text Extraction', 'PDF: PyPDF2, DOCX: docx2txt'),
        ('2. Text Cleaning', 'Remove URLs, special chars, normalize whitespace'),
        ('3. Entity Extraction', 'Skills via NER + Keyword, Experience via Regex'),
        ('4. Feature Generation', 'Sentence Embeddings, Skill Matching'),
        ('5. Model Inference', 'XGBoost Prediction, Probability Calculation')
    ]
    
    for step, detail in steps:
        p = doc.add_paragraph()
        p.add_run(f'{step}: ').bold = True
        p.add_run(detail)
    
    doc.add_paragraph()
    
    # 10. Model Training Process
    doc.add_heading('7. Model Training Process', level=1)
    
    doc.add_heading('Training Data', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Dataset 1: Career Classification\n').bold = True
    p.add_run('• Source: dataset9000.csv\n• Size: ~9,000 samples\n• Features: 17 skill dimensions\n• Target: Career role (24 classes)')
    
    p = doc.add_paragraph()
    p.add_run('Dataset 2: Job Match Prediction\n').bold = True
    p.add_run('• Source: job_dataset.csv + Synthetic Generation\n• Size: ~600 samples\n• Target: Match score (0-100)')
    
    doc.add_paragraph()
    
    # 11. Ensemble Methods
    doc.add_heading('8. Ensemble Methods', level=1)
    
    doc.add_heading('Soft-Voting Ensemble', level=2)
    p = doc.add_paragraph()
    p.add_run('The system combines predictions from multiple models:')
    
    p = doc.add_paragraph('1. Get Job Probability Model prediction')
    p.paragraph_format.left_indent = Cm(1)
    
    p = doc.add_paragraph('2. Get Career Classification prediction')
    p.paragraph_format.left_indent = Cm(1)
    
    p = doc.add_paragraph('3. Apply soft-voting bonus (max +10%)')
    p.paragraph_format.left_indent = Cm(1)
    
    p = doc.add_paragraph('4. Final probability = min(100%, job_prob + bonus)')
    p.paragraph_format.left_indent = Cm(1)
    
    doc.add_heading('Fallback Strategy', level=2)
    p = doc.add_paragraph()
    p.add_run('If ML models fail to load, a rule-based fallback calculates:')
    p = doc.add_paragraph('probability = (matched_skills / total_job_skills) * 100')
    
    doc.add_paragraph()
    
    # 12. API Endpoints
    doc.add_heading('9. API Endpoints', level=1)
    
    endpoints = [
        ('POST /analyze_resume', 'Analyze resume and predict career paths'),
        ('POST /predict', 'Predict career based on skill test responses'),
        ('POST /predict-job-probability', 'Calculate job fit for target position')
    ]
    
    for endpoint, desc in endpoints:
        p = doc.add_paragraph()
        p.add_run(f'{endpoint}: ').bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    
    # 13. Technical Stack
    doc.add_heading('10. Technical Stack', level=1)
    
    stack = [
        ('Backend', 'Flask + Flask-CORS'),
        ('ML Framework', 'scikit-learn, XGBoost'),
        ('NLP', 'spaCy, Sentence Transformers'),
        ('Text Processing', 'PyPDF2, docx2txt'),
        ('Data Processing', 'Pandas, NumPy'),
        ('Serialization', 'Pickle, Joblib')
    ]
    
    for component, tech in stack:
        p = doc.add_paragraph()
        p.add_run(f'{component}: ').bold = True
        p.add_run(tech)
    
    doc.add_paragraph()
    
    # 14. Future Improvements
    doc.add_heading('11. Future Improvements', level=1)
    
    improvements = [
        ('Deep Learning Enhancement', 'Fine-tune BERT/RoBERTa for career classification'),
        ('Data Augmentation', 'Use GPT-based data generation'),
        ('Multi-modal Analysis', 'GitHub profile, LinkedIn integration'),
        ('Real-time Learning', 'User feedback integration'),
        ('Advanced Analytics', 'Career progression, salary estimation')
    ]
    
    for area, improvement in improvements:
        p = doc.add_paragraph()
        p.add_run(f'{area}: ').bold = True
        p.add_run(improvement)
    
    doc.add_paragraph()
    
    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Document Version: 1.0 | Career Guidance System')
    run.font.size = Pt(10)
    
    # Save
    save_path = os.path.join(os.path.dirname(__file__), 'Algorithm_Documentation.docx')
    doc.save(save_path)
    print(f"Generated DOCX at: {save_path}")
    return save_path

if __name__ == "__main__":
    create_algorithm_docx()
